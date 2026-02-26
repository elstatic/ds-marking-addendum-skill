#!/usr/bin/env python3
"""Generate a marking addendum DOCX from a template and client metadata.

The script keeps template wording (section texts) and applies clean business formatting.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.shared import Cm, Pt
except ModuleNotFoundError:
    raise SystemExit(
        "Missing dependency: python-docx.\n"
        "Install with:\n"
        "  python3 -m venv .venv && source .venv/bin/activate && pip install python-docx\n"
        "Then rerun this script from the same shell."
    )

MONTHS = {
    "01": "января",
    "02": "февраля",
    "03": "марта",
    "04": "апреля",
    "05": "мая",
    "06": "июня",
    "07": "июля",
    "08": "августа",
    "09": "сентября",
    "10": "октября",
    "11": "ноября",
    "12": "декабря",
}


def clean(text: str) -> str:
    return text.replace("\u2028", "\n").replace("\xa0", " ").strip()


def format_ru_date(date_ddmmyyyy: str) -> str:
    day, month, year = date_ddmmyyyy.split(".")
    return f"«{day}» {MONTHS[month]} {year} года"


def extract_template_blocks(template_path: Path) -> dict[str, str]:
    doc = Document(str(template_path))
    p = [clean(x.text) for x in doc.paragraphs]

    # Expected structure from the working template used by the team.
    # Fallbacks are included to keep the script usable with near-identical templates.
    out: dict[str, str] = {
        "h1": p[6] if len(p) > 6 else "1. Предмет соглашения",
        "p12": p[8] if len(p) > 8 else "1.2. Услуги по сопровождению обязательной маркировки рекламы включают:",
        "bullets": p[9]
        if len(p) > 9
        else "— администрирование размещений, подлежащих обязательной маркировке;\n"
        "— взаимодействие с рекламной платформой;\n"
        "— контроль корректности передачи данных;\n"
        "— формирование и сверку отчётности по маркируемой рекламе.",
        "h2": p[10] if len(p) > 10 else "2. Стоимость услуг и порядок расчётов",
        "p21": p[11]
        if len(p) > 11
        else "2.1. Стоимость Услуг маркировки составляет 3% от объёма размещений, подлежащих обязательной маркировке (за исключением поисковых размещений), согласно официальным данным рекламной платформы.",
        "p22": p[12]
        if len(p) > 12
        else "2.2. Расчёт стоимости производится на основании отчёта платформы, формируемого в месяце, следующем за отчётным.",
        "p23": p[13]
        if len(p) > 13
        else "2.3. В случае корректировок платформы перерасчёт стоимости производится в следующем отчётном периоде.",
        "h3": p[14] if len(p) > 14 else "3. Сроки оплаты",
        "p31": p[15]
        if len(p) > 15
        else "3.1. Принципал обязуется оплатить сумму, рассчитанную по п.2.2, в течение 10 календарных дней с момента выставления Агентом счёта на оплату.",
        "p32": p[16]
        if len(p) > 16
        else "3.2. Оплата производится на расчётный счёт Агента, указанный в Договоре, с обязательной ссылкой на номер и дату настоящего Дополнительного соглашения.",
        "h4": p[17] if len(p) > 17 else "4. Прочие условия",
        "p41": p[18]
        if len(p) > 18
        else "4.1. Настоящее Дополнительное соглашение является неотъемлемой частью Договора.",
        "p42": p[19]
        if len(p) > 19
        else "4.2. Все остальные условия Договора остаются без изменений и сохраняют силу.",
        "p43": p[20]
        if len(p) > 20
        else "4.3. Настоящее Дополнительное соглашение вступает в силу с момента подписания обеими сторонами.",
        "h5": p[21] if len(p) > 21 else "5. Подписи сторон",
    }

    return out


def set_doc_defaults(doc: Document) -> None:
    sec = doc.sections[0]
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)


def add_para(
    doc: Document,
    text: str,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    bold: bool = False,
    size: int = 12,
    before: int = 0,
    after: int = 6,
    first_indent: bool = True,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.15
    if first_indent:
        pf.first_line_indent = Cm(1.25)

    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_two_col(doc: Document, left: str, right: str, *, bold: bool = False, after: int = 3) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_after = Pt(after)
    pf.line_spacing = 1.15
    pf.tab_stops.add_tab_stop(Cm(8.8), WD_TAB_ALIGNMENT.LEFT)

    l = p.add_run(left)
    p.add_run("\t")
    r = p.add_run(right)
    for run in (l, r):
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def agreement_label_dative(kind: str) -> str:
    return "агентскому договору" if kind == "agent" else "договору"


def agreement_label_genitive(kind: str) -> str:
    return "агентского договора" if kind == "agent" else "договора"


def build_doc(args: argparse.Namespace) -> Document:
    template = extract_template_blocks(Path(args.template))

    doc = Document()
    set_doc_defaults(doc)

    title = f"Дополнительное соглашение № {args.ds_no}"
    subtitle = f"к {agreement_label_dative(args.agreement_kind)} № {args.agreement_no} от {args.agreement_date}."

    add_para(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, after=6, first_indent=False)
    add_para(doc, subtitle, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=10, first_indent=False)

    city_line = doc.add_paragraph()
    city_line.paragraph_format.space_after = Pt(8)
    city_line.paragraph_format.line_spacing = 1.15
    city_line.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
    city_run = city_line.add_run(f"{args.city}\t{format_ru_date(args.sign_date)}")
    city_run.font.name = "Times New Roman"
    city_run.font.size = Pt(12)

    intro = (
        "Индивидуальный предприниматель Замятин Николай Григорьевич "
        "(ОГРНИП 313590433900045), именуемый в дальнейшем «Агент», с одной стороны, "
        f"и {args.principal_full}, в лице {args.principal_position_intro} {args.principal_signer_full}, "
        f"{args.acting_word} на основании Устава, именуемое, в дальнейшем «Принципал», "
        "с другой стороны, совместно именуемые «Стороны», заключили настоящее "
        f"Дополнительное соглашение №{args.ds_no} (далее - Соглашение) о нижеследующем:"
    )
    add_para(doc, intro, after=8)

    p11 = (
        "1.1. Настоящее Дополнительное соглашение регулирует порядок расчётов и оказания услуг по "
        "сопровождению обязательной маркировки рекламы (далее — «Услуги маркировки») в рамках "
        f"{agreement_label_genitive(args.agreement_kind)} №\u202f{args.agreement_no} "
        f"от {format_ru_date(args.agreement_date)[: -5]} г. (далее — «Договор»)."
    )

    add_para(doc, template["h1"], align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, before=8, after=4, first_indent=False)
    add_para(doc, p11, after=4, first_indent=False)
    add_para(doc, template["p12"], after=3, first_indent=False)

    items = [x.strip(" -—\t") for x in re.split(r"[\n\r]+", template["bullets"]) if x.strip()]
    for item in items:
        add_para(doc, f"• {item}", after=2, first_indent=False)

    add_para(doc, template["h2"], align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, before=8, after=4, first_indent=False)
    add_para(doc, template["p21"], after=3, first_indent=False)
    add_para(doc, template["p22"], after=3, first_indent=False)
    add_para(doc, template["p23"], after=6, first_indent=False)

    add_para(doc, template["h3"], align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, before=8, after=4, first_indent=False)
    add_para(doc, template["p31"], after=3, first_indent=False)
    add_para(doc, template["p32"], after=6, first_indent=False)

    add_para(doc, template["h4"], align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, before=8, after=4, first_indent=False)
    add_para(doc, template["p41"], after=3, first_indent=False)
    add_para(doc, template["p42"], after=3, first_indent=False)
    add_para(doc, template["p43"], after=8, first_indent=False)

    add_para(doc, template["h5"], align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, before=8, after=4, first_indent=False)
    add_two_col(doc, "Принципал", "Агент", bold=True, after=4)
    add_two_col(doc, args.principal_short, "ИП Замятин Николай Григорьевич", after=2)
    add_two_col(doc, args.principal_position_sign, "", after=8)
    add_two_col(doc, "____________ М.П.", "____________ М.П.", after=2)
    add_two_col(doc, args.principal_signer_short, "Замятин Н.Г.", after=0)

    return doc


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate marking DS DOCX from template.")
    parser.add_argument("--template", required=True, help="Path to template .docx")
    parser.add_argument("--output", required=True, help="Output .docx path")

    parser.add_argument("--ds-no", required=True, help="ДС номер (e.g. 1, 2)")
    parser.add_argument("--agreement-kind", choices=["agent", "contract"], default="agent")
    parser.add_argument("--agreement-no", required=True, help="Contract number, e.g. AG-092023-1880")
    parser.add_argument("--agreement-date", required=True, help="DD.MM.YYYY")
    parser.add_argument("--sign-date", required=True, help="DD.MM.YYYY")
    parser.add_argument("--city", default="г. Пермь")

    parser.add_argument("--principal-full", required=True, help="Full legal name in intro")
    parser.add_argument("--principal-short", required=True, help="Short name for signature block")
    parser.add_argument(
        "--principal-position-intro",
        required=True,
        help="Position in genitive case for intro, e.g. 'Генерального директора'",
    )
    parser.add_argument(
        "--principal-position-sign",
        required=True,
        help="Position for signature block, e.g. 'Генеральный директор'",
    )
    parser.add_argument("--principal-signer-full", required=True, help="FIO in genitive case")
    parser.add_argument("--principal-signer-short", required=True, help="FIO short, e.g. Иванов И.И.")
    parser.add_argument(
        "--acting-word",
        default="действующего",
        help="Word in intro: действующего / действующей",
    )

    return parser.parse_args()


def main() -> None:
    args = parse_args()
    out = Path(args.output)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = build_doc(args)
    doc.save(str(out))
    print(out)


if __name__ == "__main__":
    main()
